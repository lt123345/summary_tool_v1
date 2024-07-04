from io import BytesIO
import streamlit as st
import docx
from pandas import DataFrame
import re
from urllib.parse import quote
import copy

st.set_page_config(layout="wide")

st.title("ç¥ç§˜çš„æ•°æ®ç»Ÿè®¡å·¥å…·")

# data_doc_path = "/Users/lantian/Desktop/zhouji/2023å¹´08æœˆè´¨æ§ç®€æŠ¥.docx"
uploaded_file = st.file_uploader("é€‰æ‹©ç®€æŠ¥æ–‡ä»¶(docx)", type="docx")
if not uploaded_file:
  st.stop()

filename_match = re.match(r"(\d+)å¹´(\d+)æœˆ\.*", uploaded_file.name)
if not filename_match:
  st.error("æ— æ³•ä»æ–‡ä»¶åä¸­è§£æå¹´ä»½å’Œæœˆä»½")
  st.stop()

year = int(filename_match.group(1))
month = int(filename_match.group(2))

record_doc = docx.Document(uploaded_file)
tables = record_doc.tables


def remove_blanks(s):
  return s.replace("\n", "").replace(" ", "").replace("\u3000", "")

def is_header_text(s):
  return "æŒ‡æ ‡" in s or "é¡¹ç›®" in s or "ç§‘å®¤" in s or "ç§‘åˆ«" in s

data_tables: list[list[list[str]]] = []
for table in tables:
  data_table: list[list[str]] = []
  for row in table.rows:
    data_table.append([remove_blanks(cell.text) for cell in row.cells])
  data_tables.append(data_table)

all_keshi: set[str] = set()
for row in data_tables[1]:
  name = remove_blanks(row[0])
  if name and name != "é¡¹ç›®" and name != "ç§‘å®¤":
    all_keshi.add(name)

all_keshi_sorted = sorted(all_keshi)

col1, _, _ = st.columns(3)
with col1:
  index = 0
  if "keshi" in st.session_state and st.session_state.keshi in all_keshi:
    index = all_keshi_sorted.index(st.session_state.keshi)

  # def handle_keshi_change():
  keshi = st.selectbox("é€‰æ‹©ç§‘å®¤", all_keshi_sorted, key="keshi_selector", index=index)
  st.session_state.keshi = keshi

# åˆå¹¶å¤šåˆ—è¡¨æ ¼æˆä¸€åˆ—
for i in range(len(data_tables)):
  table = data_tables[i]
  if len(table) < 1:
    continue
  headers = table[0]
  first_header = headers[0]
  same_headers = [value for value in headers if value == first_header]
  step = int(len(headers) / len(same_headers))
  if (len(headers) != len(same_headers) and
      "åŒ»é™¢åç§°" not in first_header and
      len(same_headers) >= 2 and
      "".join(headers[0:step]) == "".join(headers[step:step+step])
  ):
    new_table: list[list[str]] = []
    for i, row in enumerate(table):
      new_table.append(row[0: step])
      if i > 0:
        for j in range(step, len(headers), step):
          new_table.append(row[j:j+step])
    data_tables[i] = new_table

def get_first_text(row):
  first_text = remove_blanks(row[0])
  if len(row) > 1:
    first_text += remove_blanks(row[1])
  if len(row) > 2:
    first_text += remove_blanks(row[2])
  return first_text

def filter_row(i: int, row: list[str]):
  first_text = get_first_text(row)
  is_header = is_header_text(first_text)
  is_keshi = i >= 1 and (keshi in row[0] or (len(row) > 1 and keshi in row[1]))
  return is_header or is_keshi

filtered_data: list[list[list[str]]] = []
for table in data_tables:
  filtered_data.append([row for i, row in enumerate(table) if filter_row(i, row)])

targets = [
  ["ç—…åºŠä½¿ç”¨ç‡",r"ç—…åºŠä½¿ç”¨ç‡"],
  ["ç—…åºŠå‘¨è½¬ç‡",r"ç—…åºŠå‘¨è½¬ç‡"],
  ["å¹³å‡ä½é™¢æ—¥",r"å¹³å‡ä½é™¢æ—¥"],
  ["æ²»æ„ˆå¥½è½¬ç‡",r"å¥½è½¬ç‡"],
  ["å…¥é™¢ä¸å‡ºé™¢è¯Šæ–­ç¬¦åˆç‡",r"å…¥é™¢ä¸å‡ºé™¢è¯Šæ–­ç¬¦åˆç‡"],
  ["æˆåˆ†è¾“è¡€ç‡",r"æˆä»½è¾“è¡€ç‡"],
  ["æŠ—èŒè¯ç‰©ä½¿ç”¨ç‡",r"æŠ—èŒè¯ç‰©ä½¿ç”¨ç‡"],
  ["æŠ—èŒè¯ç‰©ä½¿ç”¨å¼ºåº¦",r"æŠ—èŒè¯ç‰©ä½¿ç”¨å¼ºåº¦"],
  ["è¯ç‰©æ„æˆæ¯”",r"ä¸šåŠ¡æ”¶å…¥ä¸å«è€—ææ”¶å…¥è¯å æ¯”"],
  ["ä¸´åºŠè·¯å¾„å…¥å¾„ç‡",r"ä¸´åºŠè·¯å¾„å…¥å¾„ç‡"],
  ["ä¸´åºŠè·¯å¾„å®Œæˆç‡",r"ä¸´åºŠè·¯å¾„å®Œæˆç‡"],
  ["ä¸´åºŠè·¯å¾„è¦†ç›–ç‡",r"ä¸´åºŠè·¯å¾„è¦†ç›–ç‡"],
  ["é™¢å†…æ„ŸæŸ“å‘ç”Ÿç‡",r"^æ„ŸæŸ“ç‡%$"],
]
results = {}
for target, search_text in targets:
  results[target] = [target, search_text, "", ""]

found_targets = set()

def find_value_v3(table, target_row_name, target_column_name):
  if len(table) == 0:
    return ["", ""]

  headers = table[0]
  column_index = next(filter(
    lambda x: re.search(target_column_name, headers[x]),
    range(len(headers))
  ), None)
  if column_index is None:
    return ["", ""]

  found_values = []
  for row in table:
    row_name = row[0] + row[1]
    cell = row[column_index]
    if target_row_name in row_name or ("ç§‘å®¤" in row_name and "â‰¥" in cell):
      found_values.append(cell)

  if len(found_values) == 0:
    return ["", ""]
  if len(found_values) == 2:
    expected, actual = found_values
    return [actual, expected]
  return [found_values[0], ""]

for table in filtered_data:
  for target, search_text in targets:
    if target in found_targets:
      continue

    actual, expected = find_value_v3(table, target_row_name=keshi, target_column_name=search_text)
    if actual:
      found_targets.add(target)
      results[target] = [target, search_text, actual, expected]

# ç‰¹åˆ«å¤„ç†çš„æŒ‡æ ‡
# ç—…ç§æ€»ä¾‹æ•°ï¼è¡¨å•å†…æ‰€æœ‰ç–¾ç—…åç§°ä¸‹ï¼Œç—…ä¾‹æ•°çš„æ€»å’Œã€‚ä¾‹å¦‚æ”¾ç–—ç§‘æ˜¯
# å…¥å¾„ç‡ï¼å…¥å¾„æ•°/ç—…ç§æ€»ç—…ä¾‹æ•°
# å®Œæˆç‡ï¼å˜å¼‚å®Œæˆä¾‹æ•°/ç—…ç§æ€»ä¾‹æ•°ã€‚
def format_float(f):
  res = f"{f:.2f}"
  if res[-1] == "0":
    res = res[:-1]
  return res

def get_special_metrics():
  tables = [t for t in filtered_data if len(t) > 0 and "ç—…ç§åç§°" in t[0]]
  if len(tables) == 0 or len(tables[0]) < 2:
    return {}
  table = tables[0]
  total_col = 2
  rujing_col = 3
  finish_col = 5
  total = 0
  rujing = 0
  finish = 0
  data_rows = table[1:]
  for row in data_rows:
    total += int(row[total_col])
    rujing += int(row[rujing_col])
    finish += int(row[finish_col])
  return {
    "ä¸´åºŠè·¯å¾„å…¥å¾„ç‡": ["ä¸´åºŠè·¯å¾„å…¥å¾„ç‡", "= å…¥å¾„æ•°/ç—…ç§æ€»ç—…ä¾‹æ•°", format_float(rujing / total * 100), ""],
    "ä¸´åºŠè·¯å¾„å®Œæˆç‡": ["ä¸´åºŠè·¯å¾„å®Œæˆç‡", "= å˜å¼‚å®Œæˆä¾‹æ•°/ç—…ç§æ€»ä¾‹æ•°", format_float(finish / total * 100), ""],
    "é‡ç‚¹ç–¾ç—…ä¾‹æ•°": ["é‡ç‚¹ç–¾ç—…ä¾‹æ•°", "= ç—…ç§æ€»ä¾‹æ•°", format_float(total), ""]
  }

results.update(get_special_metrics())

results_df = DataFrame(results.values(), columns=["æŒ‡æ ‡", "æœç´¢æ–¹æ³•", "å®é™…å€¼", "ç›®æ ‡å€¼"])
results_df.sort_values(by="æŒ‡æ ‡")

# æ‰¾åˆ°è¯Šæ–­é—®é¢˜
def search_paragraph(search_text):
  results = []
  for paragraph in record_doc.paragraphs:
    text = remove_blanks(paragraph.text)
    if search_text in text:
      results.append(text)
  return results

wrong_diagnose = search_paragraph(f"{keshi}å¤„æ–¹å·")
wrong_diagnose = [f"ï¼ˆ{i+1}ï¼‰{text[4:]}" for i, text in enumerate(wrong_diagnose)]

# æ‰¾åˆ°ç¯èŠ‚ç—…å†
def get_bingli():
  tables = [t for t in filtered_data if len(t) > 0 and "æ‚£è€…å§“å" in t[0] and "ä½é™¢å·" in t[0]]
  if len(tables) == 0 or len(tables[0]) < 2:
    return []
  table = tables[0]
  bingli = [row[0: 5] for row in table[1:]]
  return bingli

bingli = get_bingli()


# st.write("## ç»“æœä¸‹è½½")
output_file_path = "./æ”¾å°„æ²»ç–—ç§‘2023å¹´03æœˆåŒ»ç–—è´¨é‡ä¸å®‰å…¨æ£€æŸ¥è®°å½•.docx"
output_doc = docx.Document(output_file_path)
output_tables = output_doc.tables

# å†™å…¥æŒ‡æ ‡
metrics_table = [table for table in output_tables if remove_blanks(table.rows[0].cells[0].text) == "æŒ‡æ ‡"][0]
for row in metrics_table.rows:
  metrics_name = remove_blanks(row.cells[0].text)
  if metrics_name in results:
    _, _, actual, expected = results[metrics_name]
    row.cells[1].text = actual
    row.cells[2].text = expected
  
  metrics_name = remove_blanks(row.cells[3].text)
  if metrics_name in results:
    _, _, actual, expected = results[metrics_name]
    row.cells[4].text = actual
    row.cells[5].text = expected

# å†™å…¥ç”²çº§ç¯èŠ‚ç—…ä¾‹
jiaji_bingli = [item for item in bingli if "ç”²çº§" in item[4]]
temp_table = []
bingli_table = output_tables[1].cell(0, 1).tables[0]

def copy_cell_properties(source_cell, dest_cell):
    cell_properties = source_cell._tc.get_or_add_tcPr()
    dest_cell._tc.remove(dest_cell._tc.get_or_add_tcPr())
    cell_properties = copy.copy(cell_properties)
    dest_cell._tc.append(cell_properties)

def copy_row_properties(source_row, dest_row):
  for i in range(len(source_row.cells)):
    source_cell = source_row.cells[i]
    dest_cell = dest_row.cells[i]
    copy_cell_properties(source_cell, dest_cell)

if len(jiaji_bingli) > 3:
  first_row = bingli_table.rows[0]
  for i in range(3, len(jiaji_bingli)):
    new_row = bingli_table.add_row()
    copy_row_properties(first_row, new_row)

for i, item in enumerate(jiaji_bingli):
  _, huanzhe, zhuyuanhao, problem, level = item
  row = bingli_table.rows[i+1]
  row.cells[0].text = str(i+1)
  for j in range(1, 5):
    row.cells[j].text = item[j]

# å†™å…¥è¯Šæ–­é—®é¢˜
wenti_cell = output_tables[1].cell(0, 10)
wenti_cell.text = ""
for text in wrong_diagnose:
  if remove_blanks(wenti_cell.text) == "":
    wenti_cell.text = text
  else:
    wenti_cell.add_paragraph(text)

# å†™å…¥ä¹™çº§ç—…ä¾‹
yiji_bingli = [item for item in bingli if "ä¹™çº§" in item[4]]
for i, item in enumerate(yiji_bingli):
  _, huanzhe, zhuyuanhao, problem, level = item
  text = f"ï¼ˆ{len(wrong_diagnose)+i+1}ï¼‰æ‚£è€…{huanzhe}ï¼ˆä½é™¢å·ï¼š{zhuyuanhao}ï¼‰ï¼Œå­˜åœ¨é—®é¢˜ï¼š{problem}"
  if remove_blanks(wenti_cell.text) == "":
    wenti_cell.text = text
  else:
    wenti_cell.add_paragraph(text)

# æ›¿æ¢ç§‘å®¤å’Œå¹´æœˆ
for p in output_doc.paragraphs:
  p.text = p.text.replace("æ”¾å°„æ²»ç–—ç§‘", f"{keshi}").replace(f"2023å¹´03æœˆ", f"{year:04d}å¹´{month:02d}æœˆ")

output = BytesIO()
output_doc.save(output)

st.write("## ä¸‹è½½ç»Ÿè®¡ç»“æœ")

download_filename = f"{keshi}{year:04d}å¹´{month:02d}æœˆåŒ»ç–—è´¨é‡ä¸å®‰å…¨æ£€æŸ¥è®°å½•.docx"
st.download_button(
    label=f"ğŸ’¾ {download_filename}",
    data=output.getvalue(),
    file_name=quote(download_filename),
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

### æ˜¾ç¤ºç»Ÿè®¡ç»“æœ

# st.write("## è§£æç»“æœé¢„è§ˆ")
# st.write("**æŒ‡æ ‡æœç´¢ç»“æœ**")
# col1, _ = st.columns(2)
# with col1:
#   st.table(results_df)

# wrong_diagnose_df = DataFrame(wrong_diagnose, columns=["é—¨æ€¥è¯Šå¤„æ–¹ç‚¹è¯„å…¬å¸ƒ"])
# st.write("**é—¨æ€¥è¯Šå¤„æ–¹ç‚¹è¯„å…¬å¸ƒ**")
# st.table(wrong_diagnose_df)

# bingli_df = DataFrame(bingli, columns=["ç§‘å®¤", "æ‚£è€…å§“å", "ä½é™¢å·", "å­˜åœ¨é—®é¢˜", "ç—…å†ç­‰çº§"])
# st.write("**ç¯èŠ‚ç—…å†**")
# st.table(bingli_df)

# st.write("----")
# st.write("## å‚è€ƒæ•°æ®")

# for table_data in filtered_data:
#   if len(table_data) >= 1:
#     st.table(table_data)
